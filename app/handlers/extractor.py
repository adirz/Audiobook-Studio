"""Document extraction handlers — convert source manuscripts to chapters."""

from abc import ABC, abstractmethod
from dataclasses import dataclass


@dataclass
class ExtractedChapter:
    idx: int
    title: str
    raw_text: str


@dataclass
class HeadingPreview:
    """Represents a detected heading for the user to review during range selection."""
    index: int
    title: str
    style: str         # e.g. "Heading 1", "pattern_match"
    paragraph_num: int  # position in the document
    preview: str       # first ~100 chars of content after heading


class ExtractorHandler(ABC):

    @abstractmethod
    def get_name(self) -> str: ...

    @abstractmethod
    def get_supported_extensions(self) -> list[str]: ...

    @abstractmethod
    def scan_headings(self, file_path: str) -> list[HeadingPreview]:
        """Quick scan of the document to find all headings/chapter breaks.
        Used for the range selection UI where the user picks start/end.
        """

    @abstractmethod
    def extract_chapters(self, file_path: str,
                         start_idx: int = 0,
                         end_idx: int | None = None) -> list[ExtractedChapter]:
        """Extract chapter text between the given heading indices (inclusive).
        Preserves italic markers (*text*) and blank lines.
        """

    @abstractmethod
    def get_full_text_preview(self, file_path: str,
                              start_para: int = 0,
                              count: int = 50) -> list[str]:
        """Return raw paragraph texts for scrollable preview.
        Used when the user wants to pick start/end by clicking on text
        rather than by heading.
        """


class DocxExtractorHandler(ExtractorHandler):
    """Extract chapters from .docx files."""

    import re as _re

    CHAPTER_PATTERN = _re.compile(
        r"^(Prologue|Epilogue|Appendix|Chapter\s+\d+|Part\s+\d+|\d+\s+.+)$",
        _re.IGNORECASE,
    )

    def get_name(self) -> str:
        return "Word Document (.docx)"

    def get_supported_extensions(self) -> list[str]:
        return [".docx"]

    def _load_doc(self, file_path: str):
        from docx import Document
        return Document(file_path)

    def scan_headings(self, file_path: str) -> list[HeadingPreview]:
        doc = self._load_doc(file_path)
        headings = []
        idx = 0

        for para_num, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            is_heading = (
                para.style.name.startswith("Heading")
                or self.CHAPTER_PATTERN.match(text)
            )

            if is_heading:
                # Grab preview: next non-empty paragraph
                preview = ""
                for future in doc.paragraphs[para_num + 1:]:
                    ft = future.text.strip()
                    if ft:
                        preview = ft[:120]
                        break

                headings.append(HeadingPreview(
                    index=idx,
                    title=text,
                    style=para.style.name,
                    paragraph_num=para_num,
                    preview=preview,
                ))
                idx += 1

        return headings

    def extract_chapters(self, file_path: str,
                         start_idx: int = 0,
                         end_idx: int | None = None) -> list[ExtractedChapter]:
        doc = self._load_doc(file_path)

        # First pass: identify heading positions
        headings = self.scan_headings(file_path)

        if not headings:
            # No headings found — treat entire document as one chapter
            full_text = self._extract_rich_text(doc.paragraphs)
            return [ExtractedChapter(idx=0, title="Full Text", raw_text=full_text)]

        if end_idx is None:
            end_idx = len(headings) - 1

        # Filter to selected range
        selected = [h for h in headings if start_idx <= h.index <= end_idx]
        if not selected:
            return []

        # Build paragraph ranges for each selected heading
        all_paras = doc.paragraphs
        chapters = []

        for i, heading in enumerate(selected):
            # Find end: next heading's paragraph_num, or end of doc
            if i + 1 < len(selected):
                end_para = selected[i + 1].paragraph_num
            elif heading.index < len(headings) - 1:
                # There are headings after our selection — stop at next
                next_global = [h for h in headings if h.index == heading.index + 1]
                end_para = next_global[0].paragraph_num if next_global else len(all_paras)
            else:
                end_para = len(all_paras)

            # Extract text with italic preservation
            chapter_paras = all_paras[heading.paragraph_num + 1: end_para]
            raw_text = self._extract_rich_text(chapter_paras)

            chapters.append(ExtractedChapter(
                idx=heading.index - start_idx,  # renumber from 0
                title=heading.title,
                raw_text=raw_text,
            ))

        return chapters

    def _extract_rich_text(self, paragraphs) -> str:
        """Convert paragraphs to text, preserving italics as *markers*."""
        lines = []
        for para in paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")  # preserve blank lines
                continue

            rich_text = ""
            for run in para.runs:
                if run.italic:
                    rich_text += f"*{run.text}*"
                else:
                    rich_text += run.text
            lines.append(rich_text)

        return "\n".join(lines)

    def get_full_text_preview(self, file_path: str,
                              start_para: int = 0,
                              count: int = 50) -> list[str]:
        doc = self._load_doc(file_path)
        paras = doc.paragraphs[start_para: start_para + count]
        return [p.text for p in paras]
